"""
File adapter for PDF and DOCX document ingestion.

Supports:
- PDF files (via PyMuPDF)
- DOCX files (via python-docx)
"""

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, Union

import structlog

from synaptiq.adapters.base import BaseAdapter
from synaptiq.core.exceptions import AdapterError, ValidationError
from synaptiq.core.schemas import CanonicalDocument, Segment, SourceType

logger = structlog.get_logger(__name__)


class FileAdapter(BaseAdapter):
    """
    Adapter for ingesting PDF and DOCX files.
    
    Extracts text content and segments by pages (PDF) or paragraphs (DOCX).
    Not registered with AdapterFactory since it handles file content, not URLs.
    """

    source_type = SourceType.PDF  # Will be set dynamically based on file type

    @classmethod
    def can_handle(cls, url: str) -> bool:
        """Check if this is a file path to a supported document."""
        if url.startswith(("http://", "https://")):
            return False
        path = Path(url)
        return path.suffix.lower() in (".pdf", ".docx")

    async def ingest(
        self,
        source: Union[str, bytes, io.BytesIO],
        user_id: str,
        filename: Optional[str] = None,
        title: Optional[str] = None,
    ) -> CanonicalDocument:
        """
        Ingest a PDF or DOCX file.
        
        Args:
            source: File path, bytes, or BytesIO object
            user_id: User ID for multi-tenant isolation
            filename: Original filename (required if source is bytes/BytesIO)
            title: Optional title override
            
        Returns:
            CanonicalDocument with page/paragraph segments
        """
        # Determine file type and get content
        if isinstance(source, str):
            # File path
            file_path = Path(source)
            if not file_path.exists():
                raise ValidationError(
                    message=f"File not found: {source}",
                    details={"path": source},
                )
            filename = filename or file_path.name
            file_bytes = file_path.read_bytes()
        elif isinstance(source, bytes):
            file_bytes = source
            if not filename:
                raise ValidationError(
                    message="Filename is required when source is bytes",
                    details={},
                )
        elif isinstance(source, io.BytesIO):
            file_bytes = source.read()
            source.seek(0)  # Reset for potential re-read
            if not filename:
                raise ValidationError(
                    message="Filename is required when source is BytesIO",
                    details={},
                )
        else:
            raise ValidationError(
                message=f"Unsupported source type: {type(source)}",
                details={},
            )

        file_ext = Path(filename).suffix.lower()
        
        logger.info(
            "Ingesting file",
            filename=filename,
            file_type=file_ext,
            user_id=user_id,
            size_bytes=len(file_bytes),
        )

        try:
            if file_ext == ".pdf":
                return await self._ingest_pdf(
                    file_bytes, user_id, filename, title
                )
            elif file_ext == ".docx":
                return await self._ingest_docx(
                    file_bytes, user_id, filename, title
                )
            else:
                raise ValidationError(
                    message=f"Unsupported file type: {file_ext}",
                    details={"filename": filename, "extension": file_ext},
                )

        except ValidationError:
            raise
        except Exception as e:
            logger.error("Error ingesting file", filename=filename, error=str(e))
            raise AdapterError(
                message=f"Failed to ingest file: {str(e)}",
                source_url=filename,
                adapter_type="file",
                cause=e,
            )

    async def _ingest_pdf(
        self,
        file_bytes: bytes,
        user_id: str,
        filename: str,
        title: Optional[str] = None,
    ) -> CanonicalDocument:
        """Ingest a PDF file using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise AdapterError(
                message="PyMuPDF (fitz) is required for PDF ingestion. Install with: pip install pymupdf",
                source_url=filename,
                adapter_type="file",
            )

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        segments = []
        full_text_parts = []
        char_offset = 0

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            
            if page_text.strip():
                full_text_parts.append(page_text)
                segments.append(
                    Segment(
                        text=page_text.strip(),
                        start_offset=char_offset,
                        end_offset=char_offset + len(page_text),
                        segment_type="page",
                        metadata={"page_number": page_num + 1},
                    )
                )
                char_offset += len(page_text)

        raw_content = "\n\n".join(full_text_parts)
        
        # Extract metadata
        metadata = doc.metadata or {}
        doc_title = title or metadata.get("title") or self._title_from_filename(filename)

        doc.close()

        return CanonicalDocument(
            user_id=user_id,
            source_type=SourceType.PDF,
            source_url=f"file://{filename}",
            source_title=doc_title,
            source_metadata={
                "filename": filename,
                "page_count": len(segments),
                "word_count": len(raw_content.split()),
                "author": metadata.get("author"),
                "subject": metadata.get("subject"),
                "creator": metadata.get("creator"),
                "producer": metadata.get("producer"),
            },
            raw_content=raw_content,
            content_segments=segments,
            created_at=self._parse_pdf_date(metadata.get("creationDate")),
        )

    async def _ingest_docx(
        self,
        file_bytes: bytes,
        user_id: str,
        filename: str,
        title: Optional[str] = None,
    ) -> CanonicalDocument:
        """Ingest a DOCX file using python-docx."""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise AdapterError(
                message="python-docx is required for DOCX ingestion. Install with: pip install python-docx",
                source_url=filename,
                adapter_type="file",
            )

        try:
            doc = Document(io.BytesIO(file_bytes))
        except PackageNotFoundError:
            raise ValidationError(
                message="Invalid DOCX file or file is corrupted",
                details={"filename": filename},
            )

        segments = []
        full_text_parts = []
        char_offset = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_parts.append(text)
                
                # Detect segment type based on style
                segment_type = "paragraph"
                if para.style and para.style.name:
                    style_name = para.style.name.lower()
                    if "heading" in style_name:
                        # Extract heading level
                        match = re.search(r"heading\s*(\d+)", style_name)
                        if match:
                            segment_type = f"heading_{match.group(1)}"
                        else:
                            segment_type = "heading_1"
                    elif "title" in style_name:
                        segment_type = "title"

                segments.append(
                    Segment(
                        text=text,
                        start_offset=char_offset,
                        end_offset=char_offset + len(text),
                        segment_type=segment_type,
                        metadata={"style": para.style.name if para.style else None},
                    )
                )
                char_offset += len(text) + 1  # +1 for newline

        raw_content = "\n".join(full_text_parts)
        
        # Extract metadata from core properties
        core_props = doc.core_properties
        doc_title = title or core_props.title or self._title_from_filename(filename)

        return CanonicalDocument(
            user_id=user_id,
            source_type=SourceType.DOCX,
            source_url=f"file://{filename}",
            source_title=doc_title,
            source_metadata={
                "filename": filename,
                "paragraph_count": len(segments),
                "word_count": len(raw_content.split()),
                "author": core_props.author,
                "subject": core_props.subject,
                "category": core_props.category,
                "keywords": core_props.keywords,
            },
            raw_content=raw_content,
            content_segments=segments,
            created_at=core_props.created,
        )

    def _title_from_filename(self, filename: str) -> str:
        """Extract a readable title from filename."""
        stem = Path(filename).stem
        # Replace common separators with spaces
        title = re.sub(r"[-_]", " ", stem)
        return title.title()

    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date format (D:YYYYMMDDHHmmSS)."""
        if not date_str:
            return None
        try:
            # Remove D: prefix if present
            if date_str.startswith("D:"):
                date_str = date_str[2:]
            # Parse basic format
            return datetime.strptime(date_str[:14], "%Y%m%d%H%M%S")
        except (ValueError, IndexError):
            return None

